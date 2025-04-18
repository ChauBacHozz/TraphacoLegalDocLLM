from backend.graph_database.preprocessing.preprocess_docx import (extract_text, normalize_bullets, 
                                     convert_text_list_to_tree, flatten_tree,
                                     preprocess_chunks, normalize_appendix_text_bullets,
                                     normalize_modified_text_bullets)
from backend.graph_database.db_uploader.save_doc_to_db import save_origin_doc_to_db, save_modified_doc_to_db
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from icecream import ic
from stqdm import stqdm
from tqdm import tqdm
import re

BATCH_SIZE = 50

def save_type1_origin_appendix_to_db(extracted_text, heading, doc_number, driver):
    pass

def save_type1_origin_pre_appendix_to_db(extracted_text, heading, doc_number, driver):
    """
        Hàm thực hiện tiền xử lý và sau đó lưu nội dung tiền mục lục của văn bản thông tư nghị định gốc 
        vào cơ sở dữ liệu.

        Args:
            extracted_text  : Đoạn văn bản thô đã được trích xuất từ doc gốc
            heading         : Tiêu đề của văn bản pháp luật
            doc_number      : Mã văn bản pháp luật
            driver          : Giao diện kết nối code python với cơ sở dữ liệu neo4j 
    """
    
    heading_idx = None
    # Kiểm tra xem trong 100 dòng đầu có xuất hiện từ "chương" không, nếu xuất hiện thì sẽ cắt nội dung
    # từ khi xuất hiện từ "chương" cho đến cuối
    for i, text in enumerate(extracted_text[:100]):
        if "chương" in text.lower():
            heading_idx = i
            break

    if heading_idx:
        extracted_text = extracted_text[heading_idx:]
    else:
        raise ValueError("Lỗi: Không tồn tại đề mục trong văn bản")

    # Tiền xử lý văn bản
    full_text = normalize_bullets(extracted_text)
    # Convert text list to tree base to manage content 
    tree = convert_text_list_to_tree(full_text)
    # Flatten tree into list of strings
    flattened_tree = flatten_tree(tree)
    # Split data into chunks
    chunks = [text[0] for text in flattened_tree]
    # chunks = [f"{path}: {text}" for path, text in flattened_tree]
    # Preprocess chunks
    metadata_lst = preprocess_chunks(chunks, heading, doc_number)
    # Lưu dữ liệu vào neo4j sử dụng các batch để giảm gánh nặng bộ nhớ
    batch_size = BATCH_SIZE    
    for i in stqdm(range(0, len(metadata_lst), batch_size)):
        save_origin_doc_to_db(metadata_lst[i:i+batch_size], driver)

def save_type1_modified_appendix_to_db(document, heading, doc_number, driver):
    """
        Hàm thực hiện tiền xử lý và sau đó lưu nội dung phụ lục của văn bản thông tư nghị định 
        sửa đổi vào cơ sở dữ liệu.

        Args:
            document        : Doc văn bản gốc
            heading         : Tiêu đề của văn bản pháp luật
            doc_number      : Mã văn bản pháp luật
            driver          : Giao diện kết nối code python với cơ sở dữ liệu neo4j 
    """
    # Trích xuất văn bản từ doc gốc
    extracted_text = []
    appendix_ids = []
    temp_idx = 0
    remove_idx = None
    for para in document.paragraphs:
        if para.text != "\xa0":
            extracted_text.append(para.text)
            temp_idx += 1
        if para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER and "phụ lục" in para.text.lower():
            appendix_ids.append(temp_idx)
        if para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER and "biểu mẫu" in para.text.lower():
            remove_idx = temp_idx - 1
    if remove_idx:
        full_text = extracted_text[appendix_ids[0] - 1:remove_idx]
    else:
        full_text = extracted_text[appendix_ids[0] - 1:]
    # Chuẩn hóa danh sách chỉ mục của từng phụ lục
    normalized_appendix_ids = [ids - appendix_ids[0] for ids in appendix_ids]
    chunks = normalize_appendix_text_bullets(full_text, normalized_appendix_ids)
    # Preprocess chunks, trả về danh sách các metadata
    metadata_lst = preprocess_chunks(chunks, heading, doc_number)

    # Lưu dữ liệu vào neo4j sử dụng các batch để giảm gánh nặng bộ nhớ
    batch_size = BATCH_SIZE    
    print("☑️ saving appendix data")
    for i in stqdm(range(0, len(metadata_lst), batch_size)):
        save_modified_doc_to_db(metadata_lst[i:i+batch_size], driver)

def save_type1_modified_pre_appendix_to_db(extracted_text, heading, doc_number, driver):
    """
        Hàm thực hiện tiền xử lý và sau đó lưu nội dung tiền mục lục của văn bản thông tư nghị định
        sửa đổi vào cơ sở dữ liệu.

        Args:
            extracted_text  : Đoạn văn bản thô đã được trích xuất từ doc gốc
            heading         : Tiêu đề của văn bản pháp luật
            doc_number      : Mã văn bản pháp luật
            driver          : Giao diện kết nối code python với cơ sở dữ liệu neo4j 
    """
    # Kiểm tra xem trong 100 dòng đầu có xuất hiện từ "chương" không, nếu xuất hiện thì sẽ cắt nội dung
    # từ khi xuất hiện từ "chương" cho đến cuối
    heading_idx = None
    for i, text in enumerate(extracted_text[:100]):
        if ("Chương" in text or "Mục" in text or "Điều" in text) and not text.isupper():
            heading_idx = i
            break
    if heading_idx:
        extracted_text = extracted_text[heading_idx:]
    else:
        print("ERROR")
        return
    # Tiền xử lý văn bản
    full_text = normalize_modified_text_bullets(extracted_text)
    # # Convert text list to tree base to manage content 
    tree = convert_text_list_to_tree(full_text)
    # # Flatten tree into list of strings
    flattened_tree = flatten_tree(tree)
    # # Split data into chunks
    chunks = [text[0] for text in flattened_tree]
    # # Preprocess chunks
    metadata_lst = preprocess_chunks(chunks, heading, doc_number)
    # Lưu dữ liệu vào neo4j sử dụng các batch để giảm gánh nặng bộ nhớ
    batch_size = BATCH_SIZE    
    for i in stqdm(range(0, len(metadata_lst), batch_size)):
        save_modified_doc_to_db(metadata_lst[i:i+batch_size], driver)

def save_type2_origin_appendix_to_db(document, heading, doc_number, driver):
    """
        Hàm thực hiện tiền xử lý và sau đó lưu nội dung phụ lục của văn bản luật gốc
        vào cơ sở dữ liệu.

        Args:
            document        : Doc văn bản gốc
            heading         : Tiêu đề của văn bản pháp luật
            doc_number      : Mã văn bản pháp luật
            driver          : Giao diện kết nối code python với cơ sở dữ liệu neo4j 
    """
    # Trích xuất văn bản từ doc gốc
    extracted_text = []
    appendix_ids = []
    temp_idx = 0
    remove_idx = None
    for para in document.paragraphs:
        if para.text != "\xa0":
            extracted_text.append(para.text)
            temp_idx += 1
        if para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER and "phụ lục" in para.text.lower():
            appendix_ids.append(temp_idx)
        if para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER and "biểu mẫu" in para.text.lower():
            remove_idx = temp_idx - 1
    if remove_idx:
        full_text = extracted_text[appendix_ids[0] - 1:remove_idx]
    else:
        full_text = extracted_text[appendix_ids[0] - 1:]
    # Chuẩn hóa danh sách chỉ mục của từng phụ lục
    normalized_appendix_ids = [ids - appendix_ids[0] for ids in appendix_ids]
    chunks = normalize_appendix_text_bullets(full_text, normalized_appendix_ids)
    # Preprocess chunks, trả về danh sách các metadata
    metadata_lst = preprocess_chunks(chunks, heading, doc_number)
    # Lưu dữ liệu vào neo4j sử dụng các batch để giảm gánh nặng bộ nhớ
    batch_size = BATCH_SIZE    
    for i in stqdm(range(0, len(metadata_lst), batch_size)):
        save_origin_doc_to_db(metadata_lst[i:i+batch_size], driver)

def save_type2_origin_pre_appendix_to_db(extracted_text, heading, doc_number, driver):
    """
        Hàm thực hiện tiền xử lý và sau đó lưu nội dung tiền mục lục của văn bản luật gốc 
        vào cơ sở dữ liệu.

        Args:
            extracted_text  : Đoạn văn bản thô đã được trích xuất từ doc gốc
            heading         : Tiêu đề của văn bản pháp luật
            doc_number      : Mã văn bản pháp luật
            driver          : Giao diện kết nối code python với cơ sở dữ liệu neo4j 
    """
    heading_idx = None
    # Kiểm tra xem trong 100 dòng đầu có xuất hiện từ "chương" không, nếu xuất hiện thì sẽ cắt nội dung
    # từ khi xuất hiện từ "chương" cho đến cuối
    for i, text in enumerate(extracted_text[:100]):
        if "chương" in text.lower():
            heading_idx = i
            break

    if heading_idx:
        extracted_text = extracted_text[heading_idx:]
    else:
        raise ValueError("Lỗi: Không tồn tại đề mục trong văn bản")
    # Tiền xử lý văn bản
    full_text = normalize_bullets(extracted_text)
    # Convert text list to tree base to manage content 
    tree = convert_text_list_to_tree(full_text)
    # Flatten tree into list of strings
    flattened_tree = flatten_tree(tree)
    # Split data into chunks
    chunks = [text[0] for text in flattened_tree]
    # Preprocess chunks, trả về danh sách metadata
    metadata_lst = preprocess_chunks(chunks, heading, doc_number)
    # Lưu dữ liệu vào neo4j sử dụng các batch để giảm gánh nặng bộ nhớ
    batch_size = BATCH_SIZE    
    for i in stqdm(range(0, len(metadata_lst), batch_size)):
        save_origin_doc_to_db(metadata_lst[i:i+batch_size], driver)

def save_type2_modified_appendix_to_db(document, heading, doc_number, driver):
    """
        Hàm thực hiện tiền xử lý và sau đó lưu nội dung phụ lục của văn bản luật sửa đổi
        vào cơ sở dữ liệu.

        Args:
            document        : Doc văn bản gốc
            heading         : Tiêu đề của văn bản pháp luật
            doc_number      : Mã văn bản pháp luật
            driver          : Giao diện kết nối code python với cơ sở dữ liệu neo4j 
    """
    # Trích xuất văn bản từ doc gốc
    extracted_text = []
    appendix_ids = []
    temp_idx = 0
    remove_idx = None
    for para in document.paragraphs:
        if para.text != "\xa0":
            extracted_text.append(para.text)
            temp_idx += 1
        if para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER and "phụ lục" in para.text.lower():
            appendix_ids.append(temp_idx)
        if para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER and "biểu mẫu" in para.text.lower():
            remove_idx = temp_idx - 1
    if remove_idx:
        full_text = extracted_text[appendix_ids[0] - 1:remove_idx]
    else:
        full_text = extracted_text[appendix_ids[0] - 1:]
    # Chuẩn hóa danh sách chỉ mục của từng phụ lục
    normalized_appendix_ids = [ids - appendix_ids[0] for ids in appendix_ids]
    chunks = normalize_appendix_text_bullets(full_text, normalized_appendix_ids)
    # Preprocess chunks, trả về danh sách các metadata
    metadata_lst = preprocess_chunks(chunks, heading, doc_number)
    # Lưu dữ liệu vào neo4j sử dụng các batch để giảm gánh nặng bộ nhớ
    batch_size = BATCH_SIZE    
    for i in stqdm(range(0, len(metadata_lst), batch_size)):
        save_modified_doc_to_db(metadata_lst[i:i+batch_size], driver) 

def save_type2_modified_pre_appendix_to_db(extracted_text, heading, doc_number, driver):
    """
        Hàm thực hiện tiền xử lý và sau đó lưu nội dung tiền mục lục của văn bản luật sửa đổi 
        vào cơ sở dữ liệu.

        Args:
            extracted_text  : Đoạn văn bản thô đã được trích xuất từ doc gốc
            heading         : Tiêu đề của văn bản pháp luật
            doc_number      : Mã văn bản pháp luật
            driver          : Giao diện kết nối code python với cơ sở dữ liệu neo4j 
    """
    heading_idx = None
    # Kiểm tra xem trong 100 dòng đầu có xuất hiện từ "chương" không, nếu xuất hiện thì sẽ cắt nội dung
    # từ khi xuất hiện từ "chương" cho đến cuối
    for i, text in enumerate(extracted_text[:100]):
        if ("Chương" in text or "Mục" in text or "Điều" in text) and not text.isupper():
            heading_idx = i
            break
    if heading_idx:
        modified_doc_id = re.search(r"luật dược số (.+)", "".join(extracted_text[:heading_idx]), re.IGNORECASE).group(1).strip().split(" ")[0]
        extracted_text = extracted_text[heading_idx:]
    else:
        print("ERROR")
        return
    # Tiền xử lý văn bản
    full_text = normalize_modified_text_bullets(extracted_text)
    # Convert text list to tree base to manage content 
    tree = convert_text_list_to_tree(full_text)
    # Flatten tree into list of strings
    flattened_tree = flatten_tree(tree)
    # Split data into chunks
    chunks = [text[0] for text in flattened_tree]
    # Preprocess chunks, trả về danh sách các metadata
    metadata_lst = preprocess_chunks(chunks, heading, doc_number)
    # Lưu dữ liệu vào neo4j sử dụng các batch để giảm gánh nặng bộ nhớ
    for metadata in metadata_lst:
        metadata["modified_doc_id"] = modified_doc_id
    batch_size = BATCH_SIZE    
    for i in stqdm(range(0, len(metadata_lst), batch_size)):
        save_modified_doc_to_db(metadata_lst[i:i+batch_size], driver, doc_type=2)
